from docx import Document
import os

def generate_report(extracted_data, output_format="docx"):
    doc = Document()
    doc.add_heading("Extracted Data Report", 0)
    
    for data in extracted_data:
        doc.add_heading(f"File: {data['file_name']}", level=1)
        doc.add_paragraph(data['text'])

    output_dir = "./output"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    report_path = os.path.join(output_dir, "extracted_data_report." + output_format)
    doc.save(report_path)

    return report_path
